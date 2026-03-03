"""
report_generator.py
===================
Export charts as PNG and generate a formatted Word document report
for Arc Sensor Analysis.

What this script does:
  1. Reads the cleaned CSV (from the Outputs folder).
  2. Regenerates all analysis charts and saves them as PNG files.
  3. Creates a Word document (.docx) with:
       - Title page section
       - Executive summary
       - Data statistics table
       - Arc event table
       - All charts embedded as images
  4. Saves the Word report to the OneDrive Reports folder.

Run with:
    python 02-Scripts/report_generator.py

Dependencies: pandas, numpy, plotly, kaleido, python-docx
"""

import os
import sys
from datetime import date

import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go

# python-docx — for creating Word documents
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# CONFIGURATION — adjust these paths to match your setup
# ---------------------------------------------------------------------------

# Where your cleaned CSV lives (produced by csv_cleaner.py or arc_monitor.py)
ONEDRIVE_BASE = os.path.join(
    os.path.expanduser("~"), "OneDrive", "Projects", "arc-sensor-analysis"
)
OUTPUTS_FOLDER = os.path.join(ONEDRIVE_BASE, "Outputs")
REPORTS_FOLDER = os.path.join(ONEDRIVE_BASE, "Reports")

# Expected name of the cleaned CSV in the Outputs folder.
# If you named it differently, change this.
CLEANED_CSV_NAME = "arc_m12_cleaned.csv"

# Statistical threshold multiplier for arc event detection
SIGMA_MULTIPLIER = 3

# ---------------------------------------------------------------------------
# CHART EXPORT HELPERS
# ---------------------------------------------------------------------------


def save_figure(fig: go.Figure, filename: str, folder: str) -> str:
    """
    Save a Plotly figure as a PNG file.

    Uses kaleido (installed alongside plotly) to render static images.
    Returns the full path of the saved PNG.
    """
    os.makedirs(folder, exist_ok=True)
    filepath = os.path.join(folder, filename)
    fig.write_image(filepath, width=1200, height=500, scale=2)
    print(f"  Saved chart: {filepath}")
    return filepath


def make_arc_chart(df: pd.DataFrame, threshold: float) -> go.Figure:
    """Arc intensity vs sample index with threshold line and event markers."""
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=df.index,
        y=df["arc"],
        mode="lines",
        name="Arc (µA/cm²)",
        line=dict(color="#1f77b4", width=1),
    ))
    events = df[df["is_event"]]
    if not events.empty:
        fig.add_trace(go.Scatter(
            x=events.index,
            y=events["arc"],
            mode="markers",
            name="Arc Event",
            marker=dict(color="red", size=6),
        ))
    fig.add_hline(
        y=threshold,
        line_dash="dash",
        line_color="red",
        annotation_text=f"Threshold ({threshold:.4f} µA/cm²)",
        annotation_position="top right",
    )
    fig.update_layout(
        title="Arc Intensity Along Track",
        xaxis_title="Sample index",
        yaxis_title="Arc (µA/cm²)",
        hovermode="x unified",
        height=450,
    )
    return fig


def make_height_chart(df: pd.DataFrame) -> go.Figure:
    """Track height above sea level."""
    fig = px.line(
        df,
        y="height",
        title="Track Height (m ASL)",
        labels={"index": "Sample index", "height": "Height (m ASL)"},
    )
    fig.update_layout(height=400)
    return fig


def make_distance_chart(df: pd.DataFrame) -> go.Figure:
    """Cumulative distance along track."""
    fig = px.line(
        df,
        y="distance",
        title="Distance Along Track (m)",
        labels={"index": "Sample index", "distance": "Distance (m)"},
    )
    fig.update_layout(height=400)
    return fig


def make_gps_map(df: pd.DataFrame) -> go.Figure | None:
    """GPS track coloured by arc intensity. Returns None if no GPS data."""
    if "latitude" not in df.columns or "longitude" not in df.columns:
        return None
    gps = df.dropna(subset=["latitude", "longitude"])
    if gps.empty:
        return None
    fig = px.scatter_mapbox(
        gps,
        lat="latitude",
        lon="longitude",
        color="arc",
        color_continuous_scale="RdYlGn_r",
        zoom=14,
        mapbox_style="open-street-map",
        title="GPS Track — Coloured by Arc Intensity",
        labels={"arc": "Arc (µA/cm²)"},
    )
    fig.update_layout(height=600)
    return fig


# ---------------------------------------------------------------------------
# WORD DOCUMENT HELPERS
# ---------------------------------------------------------------------------


def heading(doc: Document, text: str, level: int = 1):
    """Add a heading to the document."""
    para = doc.add_heading(text, level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return para


def body_text(doc: Document, text: str):
    """Add a normal paragraph of body text."""
    doc.add_paragraph(text)


def add_stats_table(doc: Document, df: pd.DataFrame, threshold: float):
    """
    Add a formatted table of dataset statistics to the document.
    Columns: Metric, Value, Unit.
    """
    stats = [
        ("Total samples", f"{len(df):,}", "—"),
        ("Arc — minimum", f"{df['arc'].min():.4f}", "µA/cm²"),
        ("Arc — maximum", f"{df['arc'].max():.4f}", "µA/cm²"),
        ("Arc — mean", f"{df['arc'].mean():.4f}", "µA/cm²"),
        ("Arc — std dev (σ)", f"{df['arc'].std():.4f}", "µA/cm²"),
        ("Detection threshold (mean + 3σ)", f"{threshold:.4f}", "µA/cm²"),
        ("Arc events detected", f"{df['is_event'].sum():,}", "—"),
    ]

    if "distance" in df.columns and df["distance"].notna().any():
        dist_range = df["distance"].max() - df["distance"].min()
        stats.append(("Distance range", f"{dist_range:.1f}", "m"))

    if "height" in df.columns and df["height"].notna().any():
        stats.append(("Height range", f"{df['height'].min():.1f} – {df['height'].max():.1f}", "m ASL"))

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    hdr[2].text = "Unit"
    for cell in hdr:
        run = cell.paragraphs[0].runs[0]
        run.bold = True

    # Data rows
    for metric, value, unit in stats:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value
        row[2].text = unit


def add_events_table(doc: Document, events_df: pd.DataFrame):
    """
    Add a table of detected arc events to the document.
    Shows up to 50 rows to keep the report manageable.
    """
    display = events_df.drop(columns=["is_event"], errors="ignore").head(50)

    if display.empty:
        doc.add_paragraph("No arc events were detected above the threshold.")
        return

    table = doc.add_table(rows=1, cols=len(display.columns))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    for i, col in enumerate(display.columns):
        hdr[i].text = col
        hdr[i].paragraphs[0].runs[0].bold = True

    # Data rows
    for _, row_data in display.iterrows():
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(round(val, 4)) if isinstance(val, float) else str(val)

    if len(events_df) > 50:
        doc.add_paragraph(
            f"(Table shows first 50 of {len(events_df):,} events. "
            "See the events CSV for the full list.)"
        )


def add_image(doc: Document, image_path: str, caption: str):
    """Add a PNG image to the document with a caption below."""
    if not os.path.isfile(image_path):
        doc.add_paragraph(f"[Image not found: {image_path}]")
        return
    doc.add_picture(image_path, width=Inches(6.0))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True


# ---------------------------------------------------------------------------
# MAIN REPORT BUILDER
# ---------------------------------------------------------------------------


def build_report(df: pd.DataFrame, threshold: float, chart_paths: dict) -> str:
    """
    Build the Word document and return the path where it was saved.
    """
    os.makedirs(REPORTS_FOLDER, exist_ok=True)

    doc = Document()

    # ── Title ─────────────────────────────────────────────────────────────
    title_para = doc.add_heading("Arc Analysis Report — Mileage Marker M12", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Date: {date.today().strftime('%d %B %Y')}")
    doc.add_paragraph("Location: Up line, approximately M12")
    doc.add_paragraph("Prepared by: Data Analysis Team")
    doc.add_page_break()

    # ── 1. Executive Summary ───────────────────────────────────────────────
    heading(doc, "1. Executive Summary")
    n_events = int(df["is_event"].sum())
    body_text(
        doc,
        f"A sensor survey detected {n_events:,} arc event(s) along the Up line "
        f"near mileage marker M12.  The detection threshold was set at "
        f"mean + {SIGMA_MULTIPLIER} standard deviations "
        f"({threshold:.4f} µA/cm²).  "
        f"This report summarises the findings, presents key statistics, "
        f"and provides chart evidence for engineering review."
    )

    # ── 2. Data Overview ───────────────────────────────────────────────────
    heading(doc, "2. Data Overview")
    body_text(
        doc,
        "The sensor data was loaded from a CSV file and parsed using an "
        "automated flexible parser.  Columns were mapped to standard names: "
        "Arc (µA/cm²), Distance (m), Height (m ASL), Latitude, Longitude."
    )

    # ── 3. Statistics ──────────────────────────────────────────────────────
    heading(doc, "3. Dataset Statistics")
    add_stats_table(doc, df, threshold)
    doc.add_paragraph("")  # spacer

    # ── 4. Charts ──────────────────────────────────────────────────────────
    heading(doc, "4. Charts")

    if "arc" in chart_paths:
        heading(doc, "4.1 Arc Intensity", level=2)
        body_text(
            doc,
            "The chart below shows arc intensity (µA/cm²) along the track. "
            "Red markers indicate readings above the detection threshold."
        )
        add_image(doc, chart_paths["arc"], "Figure 1 — Arc Intensity Along Track")
        doc.add_paragraph("")

    if "height" in chart_paths:
        heading(doc, "4.2 Track Height", level=2)
        body_text(doc, "Track height above sea level along the survey route.")
        add_image(doc, chart_paths["height"], "Figure 2 — Track Height (m ASL)")
        doc.add_paragraph("")

    if "distance" in chart_paths:
        heading(doc, "4.3 Distance Along Track", level=2)
        body_text(doc, "Cumulative distance measured along the track.")
        add_image(doc, chart_paths["distance"], "Figure 3 — Distance Along Track (m)")
        doc.add_paragraph("")

    if "map" in chart_paths:
        heading(doc, "4.4 GPS Track Map", level=2)
        body_text(
            doc,
            "The GPS track is coloured by arc intensity "
            "(red = high, green = low)."
        )
        add_image(doc, chart_paths["map"], "Figure 4 — GPS Track Coloured by Arc Intensity")
        doc.add_paragraph("")

    # ── 5. Detected Events ────────────────────────────────────────────────
    heading(doc, "5. Detected Arc Events")
    events_df = df[df["is_event"]].copy()
    body_text(
        doc,
        f"{n_events:,} arc event(s) were detected above the threshold of "
        f"{threshold:.4f} µA/cm².  "
        f"The table below lists the event details."
    )
    add_events_table(doc, events_df)

    # ── 6. Conclusion ─────────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "6. Conclusion and Recommended Actions")
    if n_events > 0:
        body_text(
            doc,
            f"The survey identified {n_events:,} arc event(s) at M12.  "
            "These readings exceed the statistical detection threshold and "
            "warrant further investigation.  "
            "Recommended actions:"
        )
        doc.add_paragraph("Visual inspection of overhead line equipment at M12.", style="List Bullet")
        doc.add_paragraph("Review maintenance records for this section.", style="List Bullet")
        doc.add_paragraph("Consider repeat survey after any corrective work.", style="List Bullet")
    else:
        body_text(
            doc,
            "No arc events were detected above the threshold.  "
            "The section appears within normal operating parameters."
        )

    # ── Save ──────────────────────────────────────────────────────────────
    report_path = os.path.join(
        REPORTS_FOLDER,
        f"Arc_Analysis_M12_{date.today().strftime('%Y%m%d')}.docx"
    )
    doc.save(report_path)
    return report_path


# ---------------------------------------------------------------------------
# MAIN ENTRY POINT
# ---------------------------------------------------------------------------


def main():
    print("=" * 60)
    print("  Arc Sensor Analysis — Report Generator")
    print("=" * 60)

    # --- Load cleaned CSV ---
    csv_path = os.path.join(OUTPUTS_FOLDER, CLEANED_CSV_NAME)
    if not os.path.isfile(csv_path):
        print(f"\nERROR: Cleaned CSV not found: {csv_path}")
        print("Please run csv_cleaner.py or export from arc_monitor.py first.")
        sys.exit(1)

    print(f"\nLoading: {csv_path}")
    df = pd.read_csv(csv_path, low_memory=False)

    # --- Detect arc column ---
    arc_col = next(
        (c for c in df.columns if "arc" in c.lower()), None
    )
    if arc_col is None:
        print("ERROR: Could not find an 'arc' column in the cleaned CSV.")
        sys.exit(1)

    # Rename to standard name if needed
    if arc_col != "arc":
        df = df.rename(columns={arc_col: "arc"})

    df["arc"] = pd.to_numeric(df["arc"], errors="coerce")

    # --- Calculate threshold ---
    arc_mean = df["arc"].mean()
    arc_std = df["arc"].std()
    threshold = arc_mean + SIGMA_MULTIPLIER * arc_std
    print(f"Arc mean: {arc_mean:.4f} µA/cm²")
    print(f"Arc std : {arc_std:.4f} µA/cm²")
    print(f"Threshold (mean + 3σ): {threshold:.4f} µA/cm²")

    # --- Mark events ---
    df["is_event"] = df["arc"] > threshold
    n_events = int(df["is_event"].sum())
    print(f"Events detected: {n_events:,}")

    # --- Convert other columns to numeric ---
    for col in ["distance", "height", "latitude", "longitude"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")

    # --- Export charts as PNG ---
    print("\nExporting charts …")
    chart_paths = {}

    fig_arc = make_arc_chart(df, threshold)
    chart_paths["arc"] = save_figure(fig_arc, "chart_arc.png", OUTPUTS_FOLDER)

    if "height" in df.columns and df["height"].notna().any():
        fig_h = make_height_chart(df)
        chart_paths["height"] = save_figure(fig_h, "chart_height.png", OUTPUTS_FOLDER)

    if "distance" in df.columns and df["distance"].notna().any():
        fig_d = make_distance_chart(df)
        chart_paths["distance"] = save_figure(fig_d, "chart_distance.png", OUTPUTS_FOLDER)

    fig_map = make_gps_map(df)
    if fig_map is not None:
        chart_paths["map"] = save_figure(fig_map, "chart_gps_map.png", OUTPUTS_FOLDER)

    # --- Build Word report ---
    print("\nBuilding Word report …")
    report_path = build_report(df, threshold, chart_paths)

    print("\n" + "=" * 60)
    print("  Report generation complete!")
    print(f"  Report: {report_path}")
    print(f"  Charts: {OUTPUTS_FOLDER}")
    print("=" * 60)


if __name__ == "__main__":
    main()
